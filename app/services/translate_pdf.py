"""
Complete workflow to translate a Japanese patent PDF to Traditional Chinese

Usage:
    python scripts/translate_pdf.py path/to/patent.pdf --output result.docx
"""

import sys
import os
from pathlib import Path
import argparse
import json
from typing import Dict, List
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from app.utils.document_processor import DocumentProcessor, PatentSectionParser
from app.database import SessionLocal
from app.models import TranslationRequest
from app.services.translation_service import TranslationService

# For output
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class PDFTranslator:
    """Complete PDF translation workflow"""
    
    def __init__(self, use_api: bool = False):
        self.doc_processor = DocumentProcessor()
        self.section_parser = PatentSectionParser()
        self.use_api = use_api
        
        if not use_api:
            self.db = SessionLocal()
            self.translation_service = TranslationService()
    
    def translate_pdf(
        self,
        pdf_path: str,
        output_path: str,
        domain: str = "semiconductor"
    ) -> Dict:
        """
        Complete translation workflow
        
        Returns:
            Statistics and results
        """
        logger.info(f"Starting translation: {pdf_path}")
        
        # 1. Extract text from PDF
        logger.info("Step 1: Extracting text from PDF...")
        text = self.doc_processor.extract_text(pdf_path)
        
        if not text:
            raise Exception("Failed to extract text from PDF")
        
        text = self.doc_processor.clean_text(text)
        logger.info(f"  Extracted {len(text)} characters")
        
        # 2. Parse into sections
        logger.info("Step 2: Parsing sections...")
        sections = self.section_parser.parse_sections(text, 'japanese')
        logger.info(f"  Found {len(sections)} sections")
        
        # 3. Translate each section
        logger.info("Step 3: Translating sections...")
        translations = {}
        stats = {
            'total_sections': len(sections),
            'translated': 0,
            'failed': 0,
            'total_examples_used': 0,
            'total_terms_matched': 0
        }
        
        for i, (section_name, section_text) in enumerate(sections.items(), 1):
            logger.info(f"  [{i}/{len(sections)}] Translating {section_name}...")
            
            try:
                result = self._translate_section(
                    section_text,
                    section_type=section_name,
                    domain=domain
                )
                
                translations[section_name] = result
                stats['translated'] += 1
                stats['total_examples_used'] += result.get('examples_used', 0)
                stats['total_terms_matched'] += result.get('terms_matched', 0)
                
                logger.info(f"    ✓ Done (confidence: {result['confidence']:.2f})")
                
            except Exception as e:
                logger.error(f"    ✗ Failed: {e}")
                translations[section_name] = {
                    'translation': f"[翻譯失敗: {str(e)}]",
                    'confidence': 0,
                    'error': str(e)
                }
                stats['failed'] += 1
        
        # 4. Generate output document
        logger.info("Step 4: Generating output document...")
        self._create_docx(translations, sections, output_path, stats)
        
        logger.info(f"✓ Translation completed: {output_path}")
        
        return {
            'input': pdf_path,
            'output': output_path,
            'stats': stats,
            'timestamp': datetime.now().isoformat()
        }
    
    def _translate_section(
        self,
        text: str,
        section_type: str,
        domain: str
    ) -> Dict:
        """Translate a single section"""
        
        # Limit section length (to avoid token limits)
        if len(text) > 5000:
            text = text[:5000] + "\n...[內容過長，已截斷]"
        
        if self.use_api:
            # Use API endpoint (for production)
            import requests
            response = requests.post(
                "http://localhost:8000/api/v1/translate",
                json={
                    "japanese_text": text,
                    "section_type": section_type,
                    "domain": domain,
                    "use_rag": True,
                    "num_examples": 3
                }
            )
            
            if response.status_code == 200:
                data = response.json()
                return {
                    'translation': data['translation'],
                    'confidence': data['confidence_score'],
                    'examples_used': data['metadata']['examples_used'],
                    'terms_matched': data['metadata']['terms_matched']
                }
            else:
                raise Exception(f"API error: {response.status_code}")
        
        else:
            # Direct service call (faster for local)
            request = TranslationRequest(
                japanese_text=text,
                section_type=section_type,
                domain=domain,
                use_rag=True,
                num_examples=3
            )
            
            response = self.translation_service.translate(request, self.db)
            
            return {
                'translation': response.translation,
                'confidence': response.confidence_score,
                'examples_used': response.metadata['examples_used'],
                'terms_matched': response.metadata['terms_matched']
            }
    
    def _create_docx(
        self,
        translations: Dict,
        original_sections: Dict,
        output_path: str,
        stats: Dict
    ):
        """Create formatted DOCX output"""
        
        doc = Document()
        
        # Title
        title = doc.add_heading('專利翻譯結果', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Statistics
        doc.add_heading('翻譯統計', 1)
        stats_para = doc.add_paragraph()
        stats_para.add_run(f"總段落數: {stats['total_sections']}\n")
        stats_para.add_run(f"成功翻譯: {stats['translated']}\n")
        stats_para.add_run(f"失敗: {stats['failed']}\n")
        stats_para.add_run(f"使用案例數: {stats['total_examples_used']}\n")
        stats_para.add_run(f"匹配術語數: {stats['total_terms_matched']}\n")
        
        doc.add_page_break()
        
        # Section order (preferred display order)
        section_order = [
            'title',
            'abstract',
            'technical_field',
            'background',
            'prior_art',
            'summary',
            'problem',
            'solution',
            'effects',
            'description',
            'examples',
            'claims',
        ]
        
        # Add remaining sections
        for key in translations.keys():
            if key not in section_order:
                section_order.append(key)
        
        # Write sections
        for section_name in section_order:
            if section_name not in translations:
                continue
            
            result = translations[section_name]
            
            # Section heading
            heading = doc.add_heading(self._format_section_name(section_name), 2)
            
            # Translation
            trans_para = doc.add_paragraph()
            trans_run = trans_para.add_run(result['translation'])
            
            # Mark low confidence in red
            if result.get('confidence', 1) < 0.5:
                trans_run.font.color.rgb = RGBColor(255, 0, 0)
            
            # Add metadata
            if not result.get('error'):
                meta_para = doc.add_paragraph()
                meta_run = meta_para.add_run(
                    f"[信心度: {result['confidence']:.2f}, "
                    f"使用案例: {result.get('examples_used', 0)}, "
                    f"術語: {result.get('terms_matched', 0)}]"
                )
                meta_run.font.size = Pt(8)
                meta_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # Spacing
        
        # Save
        doc.save(output_path)
    
    def _format_section_name(self, name: str) -> str:
        """Format section name for display"""
        mapping = {
            'title': '發明名稱',
            'abstract': '摘要',
            'technical_field': '技術領域',
            'background': '背景技術',
            'prior_art': '先前技術',
            'summary': '發明內容',
            'problem': '欲解決之問題',
            'solution': '解決問題之技術手段',
            'effects': '發明之功效',
            'description': '實施方式',
            'examples': '實施例',
            'claims': '申請專利範圍',
        }
        
        # Handle para_XXXX
        if name.startswith('para_'):
            return f"段落 {name.replace('para_', '')}"
        
        # Handle claim_X
        if name.startswith('claim_'):
            return f"請求項 {name.replace('claim_', '')}"
        
        return mapping.get(name, name)
    
    def __del__(self):
        if hasattr(self, 'db'):
            self.db.close()


def main():
    parser = argparse.ArgumentParser(
        description='Translate Japanese patent PDF to Traditional Chinese'
    )
    parser.add_argument('pdf_path', help='Input PDF file')
    parser.add_argument('--output', '-o', help='Output DOCX file',
                       default='translation_output.docx')
    parser.add_argument('--domain', '-d', 
                       choices=['semiconductor', 'mechanical', 'general'],
                       default='semiconductor',
                       help='Patent domain')
    parser.add_argument('--api', action='store_true',
                       help='Use API endpoint instead of direct service')
    parser.add_argument('--report', '-r',
                       help='Save translation report as JSON')
    
    args = parser.parse_args()
    
    # Check input file
    if not Path(args.pdf_path).exists():
        print(f"Error: File not found: {args.pdf_path}")
        sys.exit(1)
    
    print("=" * 80)
    print("Patent Translation System")
    print("=" * 80)
    print(f"\nInput:  {args.pdf_path}")
    print(f"Output: {args.output}")
    print(f"Domain: {args.domain}")
    print()
    
    # Translate
    try:
        translator = PDFTranslator(use_api=args.api)
        
        result = translator.translate_pdf(
            args.pdf_path,
            args.output,
            domain=args.domain
        )
        
        print("\n" + "=" * 80)
        print("Translation Summary")
        print("=" * 80)
        print(f"Total sections:    {result['stats']['total_sections']}")
        print(f"Translated:        {result['stats']['translated']}")
        print(f"Failed:            {result['stats']['failed']}")
        print(f"Examples used:     {result['stats']['total_examples_used']}")
        print(f"Terms matched:     {result['stats']['total_terms_matched']}")
        print("=" * 80)
        
        # Save report
        if args.report:
            with open(args.report, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            print(f"\nReport saved: {args.report}")
        
        print(f"\n✓ Translation completed successfully!")
        print(f"Output: {args.output}")
        
    except Exception as e:
        print(f"\n✗ Translation failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()


# ===== scripts/batch_translate.py =====
"""
Batch translate multiple PDF files
"""

import sys
from pathlib import Path
import argparse
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

def batch_translate(input_dir: str, output_dir: str, domain: str):
    """Translate all PDFs in a directory"""
    
    input_path = Path(input_dir)
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)
    
    # Find all PDFs
    pdf_files = list(input_path.glob("*.pdf"))
    
    print(f"Found {len(pdf_files)} PDF files")
    
    # Import here to avoid circular import
    from scripts.translate_pdf import PDFTranslator
    
    translator = PDFTranslator()
    
    results = []
    
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n[{i}/{len(pdf_files)}] Processing: {pdf_file.name}")
        
        output_file = output_path / f"{pdf_file.stem}_translated.docx"
        
        try:
            result = translator.translate_pdf(
                str(pdf_file),
                str(output_file),
                domain=domain
            )
            results.append({
                'file': pdf_file.name,
                'status': 'success',
                'output': str(output_file),
                'stats': result['stats']
            })
            
        except Exception as e:
            print(f"  ✗ Failed: {e}")
            results.append({
                'file': pdf_file.name,
                'status': 'failed',
                'error': str(e)
            })
    
    # Summary
    print("\n" + "=" * 80)
    print("Batch Translation Summary")
    print("=" * 80)
    
    successful = sum(1 for r in results if r['status'] == 'success')
    failed = sum(1 for r in results if r['status'] == 'failed')
    
    print(f"Total files:    {len(pdf_files)}")
    print(f"Successful:     {successful}")
    print(f"Failed:         {failed}")
    
    # Save report
    import json
    report_file = output_path / f"batch_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    with open(report_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    
    print(f"\nReport saved: {report_file}")


def main():
    parser = argparse.ArgumentParser(description='Batch translate PDFs')
    parser.add_argument('input_dir', help='Directory with PDF files')
    parser.add_argument('output_dir', help='Output directory')
    parser.add_argument('--domain', '-d', default='semiconductor')
    
    args = parser.parse_args()
    
    batch_translate(args.input_dir, args.output_dir, args.domain)


if __name__ == "__main__":
    main()